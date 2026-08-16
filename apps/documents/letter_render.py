"""Rendering eines Brief-Entwurfs ins Vorlagen-Layout (#1095): **Word**
(.docx) als editierbarer Master und **PDF** als Druck-/Sende-Artefakt.

Beide Ausgaben entstehen aus derselben `LetterContent` -- dem einmal
zusammengebauten Brief aus Briefkopf, Absender-/Empfängerblock,
Datumszeile, Betreff, Text, Grußformel und Signatur. Genau deshalb gibt es
diese Zwischenstufe: zwei Renderer, die sich ihre Bestandteile jeweils
selbst aus dem Entwurf zusammensuchen, wären zwei Gelegenheiten, das
Layout unterschiedlich zu interpretieren.

**PDF direkt gerendert (fpdf2), nicht aus dem docx konvertiert.** Die
Alternative -- LibreOffice headless als Konverter -- bräuchte ein ~400 MB
schweres Office-Paket im Container samt Subprozess-Aufruf, Timeout- und
Profil-Handling, nur um ein Layout zu reproduzieren, das wir hier ohnehin
selbst beschreiben. Der Preis für den direkten Weg ist, dass beide
Renderer gepflegt werden müssen; das ist überschaubar, weil sie aus
derselben Struktur lesen und die Layout-Entscheidungen (Ränder,
Positionen) hier oben als Konstanten stehen. fpdf2 ist zudem schon für die
Mail-Body-PDFs im Einsatz (#1070), es kommt also keine Abhängigkeit dazu.

Layout-Grundlage ist der deutsche Geschäftsbrief (DIN-5008-nah, siehe
`DEFAULT_LETTER_LAYOUT`): Anschriftfeld links oben im festen Abstand,
Absender-Kurzzeile darüber, Datumszeile rechts, Betreff fett, danach der
Text. Die Werte sind Millimeter-Konstanten statt einer Layout-Engine --
mehr braucht ein einseitiger Geschäftsbrief nicht, und alles, was eine
Vorlage daran ändern darf, kommt aus ihrem `layout`-JSON.

**Word folgt DIN 5008 Form B über Textrahmen, nicht über Fließtext**
(#1151): Anschriftfeld, Informationsblock und Falz-/Lochmarken sind
`w:framePr`-Absätze, an der Seite statt am Textfluss verankert. Eine
Anschrift mit einer Zeile mehr oder weniger (Zusatzvermerk, zweizeiliger
Firmenname) verschiebt dadurch weder sich selbst noch die Betreffzeile --
der einzige Fließtext-Absatz, der Vorschub braucht (`SUBJECT_LINE_TOP_MM`
als `space_before`), ist unabhängig von allem, was in Rahmen davor steht.
Die eigene Identität erscheint zusätzlich als Fußzeile
(`section.footer`/`first_page_footer`, weil `titlePg` Kopf und Fuß der
ersten Seite getrennt schaltet); die Kaufmanns-Zeile darin nur, wenn der
Absender-Kontakt `is_own_business` ist. Der PDF-Renderer bleibt bewusst
beim einfacheren Fließtext-Ansatz mit absoluter Y-Position (siehe unten) --
das Ticket, das die Word-Ausgabe auf Form B gebracht hat, schließt einen
PDF-Export ausdrücklich aus.
"""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass
from io import BytesIO

from django.utils.formats import date_format

logger = logging.getLogger(__name__)

# Seitenmaße/Ränder in Millimetern (A4, DIN-5008-Form-B-nah).
PAGE_WIDTH_MM = 210
PAGE_HEIGHT_MM = 297
PAGE_MARGIN_LEFT_MM = 25
PAGE_MARGIN_RIGHT_MM = 20
PAGE_MARGIN_TOP_MM = 20
PAGE_MARGIN_BOTTOM_MM = 20

# Anschriftfeld: Zusatz-/Vermerkzone (Rücksendeangabe) oben, Anschriftzone
# darunter, zusammen 40 mm hoch -- der Block landet damit im Sichtfenster
# eines Fensterumschlags DIN lang, wenn der Brief normgerecht gefaltet wird.
ADDRESS_FIELD_TOP_MM = 45
ADDRESS_FIELD_WIDTH_MM = 85
ADDRESS_FIELD_HEIGHT_MM = 40

# Informationsblock rechts (Datum): auf Höhe der Anschriftzone, in der
# rechten Spalte neben dem Anschriftfeld -- Standard-Referenzlinie nach
# DIN 5008 Form B.
INFO_BLOCK_TOP_MM = 50.8
INFO_BLOCK_WIDTH_MM = 65

# Betreffzeile: fester Abstand unter dem Anschriftfeld (DIN-5008-Referenz),
# nicht aus der Höhe des Briefkopfs abgeleitet -- ein Textrahmen bleibt an
# Ort und Stelle, egal wie hoch der Kopf ausfällt.
SUBJECT_LINE_TOP_MM = 100.6

# Falz- und Lochmarken am linken Rand, als feine Linie.
FOLD_MARK_1_MM = 87
FOLD_MARK_2_MM = 192
PUNCH_MARK_MM = 148.5
_REGISTRATION_MARK_WIDTH_MM = 5

LOGO_WIDTH_MM = 40

_BODY_FONT_SIZE = 11
_SMALL_FONT_SIZE = 8
_FOOTER_FONT_SIZE = 7


@dataclass(frozen=True)
class LetterContent:
    """Der fertig zusammengesetzte Brief -- alles, was auf das Papier
    kommt, in der Reihenfolge, in der es dort steht.

    Reine Daten, keine Model-Referenz: die Renderer sollen nicht wissen,
    ob der Brief aus einem `LetterDraft`, einer Vorschau oder einem Test
    kommt.
    """

    subject: str = ""
    body: str = ""
    letterhead: str = ""
    sender_line: str = ""
    sender_block: str = ""
    recipient_block: str = ""
    date_line: str = ""
    closing: str = ""
    signature: str = ""
    logo: bytes | None = None
    logo_name: str = ""
    sender_footer_lines: tuple[str, ...] = ()


def _lines(text: str) -> list[str]:
    return [line.strip() for line in (text or "").splitlines() if line.strip()]


def _sender_line(sender_block: str) -> str:
    """Absender-Kurzzeile über dem Anschriftfeld: der mehrzeilige
    Absenderblock in einer Zeile, durch „· " getrennt.

    DIN 5008 sieht sie klein und einzeilig vor -- sie ist die Rücksendeangabe
    im Sichtfenster, nicht der Absenderblock selbst.
    """
    return " · ".join(_lines(sender_block))


def _address_line(name: str, address: str) -> str:
    parts = [name.strip()] if (name or "").strip() else []
    parts.extend(_lines(address))
    return " · ".join(parts)


def _sender_footer_lines(sender) -> list[str]:
    """Fußzeilentext der eigenen Identität (#1151): Name/Adresse/Kontakt in
    einer Zeile, kaufmännische Angaben in einer zweiten -- unabhängig vom
    Anschriftfeld-Umschalter `show_sender_block`, weil die Fußzeile ein
    eigenständiges, immer sichtbares Element ist, nicht Teil der
    Rücksendeangabe im Anschriftfeld.

    Nur ausgeben, was hinterlegt ist: ein fehlender Wert entfällt samt
    Trennzeichen, nie eine Zeile wie „Telefon: -". Die kaufmännischen
    Angaben (USt-IdNr., Steuernummer, IBAN) erscheinen nur, wenn der
    Absender als eigenes Gewerbe markiert ist (`is_own_business`) -- ein
    privater Brief bekommt keine Steuerdaten in die Fußzeile.
    """
    if sender is None:
        return []
    contact = [
        part
        for part in (
            _address_line(sender.name, sender.address),
            (sender.email or "").strip(),
            (sender.phone or "").strip(),
        )
        if part
    ]
    lines = [" · ".join(contact)] if contact else []
    if sender.is_own_business:
        business = [
            part
            for part in (
                f"USt-IdNr. {sender.vat_id.strip()}" if (sender.vat_id or "").strip() else "",
                f"St-Nr. {sender.tax_number.strip()}" if (sender.tax_number or "").strip() else "",
                f"IBAN {sender.iban.strip()}" if (sender.iban or "").strip() else "",
            )
            if part
        ]
        if business:
            lines.append(" · ".join(business))
    return lines


# Grußformeln, wie sie am Anfang einer Signatur stehen könnten, wenn eine
# Vorlage sie versehentlich zusätzlich zum Layout-Feld `closing` einträgt
# (#1151) -- dieselbe Formelliste wie `letter_generation._CLOSING_ECHO_RE`,
# hier aber gegen den *Anfang* der Signatur statt gegen das *Ende* eines
# KI-Texts geprüft, deshalb eine eigene, kleinere Regel statt geteiltem Code.
_CLOSING_PHRASES_RE = re.compile(
    r"(mit freundlichen gr(ü|ue)(ß|ss)en|"
    r"freundliche gr(ü|ue)(ß|ss)e|hochachtungsvoll|beste gr(ü|ue)(ß|ss)e|"
    r"viele gr(ü|ue)(ß|ss)e)"
)


def _strip_duplicate_closing(signature: str) -> str:
    """Die Grußformel darf nur einmal erscheinen (Layout-Feld `closing`).

    Trägt eine Vorlage sie zusätzlich am Anfang ihrer Signatur ein -- der
    real beobachtete Fehler --, entfernt das hier die führende Zeile samt
    einer eventuell folgenden Leerzeile, statt sie ein zweites Mal neben
    `closing` zu drucken.
    """
    text = (signature or "").strip("\n")
    if not text:
        return ""
    lines = text.split("\n")
    first = lines[0].strip().rstrip(",").lower()
    if _CLOSING_PHRASES_RE.fullmatch(first):
        lines = lines[1:]
        while lines and not lines[0].strip():
            lines.pop(0)
    return "\n".join(lines).strip()


def build_content(draft) -> LetterContent:
    """`LetterDraft` -> `LetterContent`, ausschließlich aus dem Snapshot.

    Liest bewusst nichts aus `draft.template` nach: das Layout wurde beim
    Anlegen eingefroren (siehe `LetterDraft`), und ein Brief darf sich
    nicht ändern, weil jemand die Vorlage bearbeitet hat. Nur das Logo
    hängt noch an der Vorlage -- eine Bilddatei je Entwurf zu kopieren
    wäre Speicher für einen Fall (Logo-Wechsel zwischen Entwurf und
    Freigabe), den niemand hat.
    """
    letterhead = draft.layout_value("letterhead") or ""
    sender_block = draft.sender_block if draft.layout_value("show_sender_block") else ""
    recipient_block = (
        draft.recipient_block if draft.layout_value("show_recipient_block") else ""
    )

    date_line = ""
    if draft.layout_value("show_date") and draft.letter_date:
        place = (draft.layout_value("date_place") or "").strip()
        formatted = date_format(draft.letter_date, "DATE_FORMAT")
        date_line = f"{place}, {formatted}" if place else formatted

    logo_bytes, logo_name = _logo_bytes(draft)

    return LetterContent(
        subject=draft.subject if draft.layout_value("show_subject") else "",
        body=draft.body_text or "",
        letterhead=letterhead,
        sender_line=_sender_line(sender_block),
        sender_block=sender_block,
        recipient_block=recipient_block,
        date_line=date_line,
        closing=(draft.layout_value("closing") or "").strip(),
        signature=_strip_duplicate_closing(draft.signature or ""),
        logo=logo_bytes,
        logo_name=logo_name,
        sender_footer_lines=tuple(_sender_footer_lines(draft.sender)),
    )


def _logo_bytes(draft) -> tuple[bytes | None, str]:
    """Das Vorlagen-Logo als Bytes, oder `(None, "")`.

    Defensiv: das Logo liegt im Object-Storage (S3/MinIO), und ein Brief
    darf nicht daran scheitern, dass die Datei dort weg oder der Bucket
    kurz nicht erreichbar ist -- dann kommt er eben ohne Logo.
    """
    template = draft.template
    if template is None or not template.logo:
        return None, ""
    try:
        with template.logo.open("rb") as handle:
            return handle.read(), template.logo.name.rsplit("/", 1)[-1]
    except Exception:
        logger.warning(
            "Brief-Entwurf %s: Logo der Vorlage %s konnte nicht geladen werden",
            draft.pk,
            template.pk,
            exc_info=True,
        )
        return None, ""


def _content_blocks(content: LetterContent) -> list[str]:
    """Betreff/Text/Grußformel/Signatur als Absatzfolge -- die Reihenfolge,
    die beide Renderer und der Klartext teilen.
    """
    blocks = []
    if content.subject.strip():
        blocks.append(content.subject.strip())
    body = (content.body or "").strip()
    if body:
        blocks.extend(paragraph.strip() for paragraph in re.split(r"\n\s*\n", body))
    if content.closing:
        blocks.append(content.closing)
    if content.signature.strip():
        blocks.append(content.signature.strip())
    return [block for block in blocks if block]


def plain_text(content: LetterContent) -> str:
    """Der Brief als Klartext -- was als `text_content` in das abgelegte
    Dokument geht und damit in Index und Suche landet.

    Enthält Kopf und Anschrift: nach einem Brief wird genauso oft über den
    Empfänger gesucht wie über den Text.
    """
    parts = [
        content.letterhead,
        content.sender_block,
        content.recipient_block,
        content.date_line,
    ]
    parts.extend(_content_blocks(content))
    return "\n\n".join(part.strip() for part in parts if part and part.strip())


# -- Word ------------------------------------------------------------------
#
# Anschriftfeld, Informationsblock und Falz-/Lochmarken sitzen als
# klassische Word-Textrahmen (`w:framePr`, verankert an der Seite) an einer
# festen Position -- unabhängig davon, wie viele Zeilen der Text davor oder
# danach hat. Nur der Fließtext (ab der Betreffzeile) ist normaler,
# mehrseitentauglicher Textfluss; ein Rahmen nimmt daran nicht teil, wird
# also von Word bei der Positionierung des Fließtexts ignoriert.


def _twips(mm_value) -> str:
    from docx.shared import Mm

    return str(Mm(mm_value).twips)


def _pin_to_page(paragraph, *, x_mm, y_mm, width_mm=None, height_mm=None, height_rule="atLeast"):
    """Verankert einen Absatz als Textrahmen an einer festen Seitenposition.

    `atLeast` statt `exact` für Inhalte mit Text: eine ungewöhnlich lange
    Anschrift oder ein langer Vorlagenname darf den Rahmen nach unten
    sprengen, statt abgeschnitten zu werden.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    pPr = paragraph._p.get_or_add_pPr()
    frame = OxmlElement("w:framePr")
    if width_mm is not None:
        frame.set(qn("w:w"), _twips(width_mm))
    if height_mm is not None:
        frame.set(qn("w:h"), _twips(height_mm))
        frame.set(qn("w:hRule"), height_rule)
    frame.set(qn("w:hAnchor"), "page")
    frame.set(qn("w:vAnchor"), "page")
    frame.set(qn("w:x"), _twips(x_mm))
    frame.set(qn("w:y"), _twips(y_mm))
    frame.set(qn("w:wrap"), "none")
    pPr.insert(0, frame)


def _add_border(paragraph, *, side, size_pt=0.5, color="999999", space_pt=1):
    """Feine Linie an einer Absatzkante -- für Falz-/Lochmarken und die
    Trennlinie über der Fußzeile.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    pPr = paragraph._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    border = OxmlElement(f"w:{side}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), str(int(size_pt * 8)))
    border.set(qn("w:space"), str(space_pt))
    border.set(qn("w:color"), color)
    borders.append(border)
    pPr.append(borders)


def _add_page_number_field(paragraph, *, size_pt=None):
    """Ein `PAGE`-Feld -- die Seitenzahl im Kopfbereich ab Seite 2 muss sich
    mit der tatsächlichen Seite mitzählen, ein fester Text käme nicht in
    Frage.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Pt

    run = paragraph.add_run()
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)


def _set_document_language(document, lang="de-DE"):
    """Dokumentsprache Deutsch statt des python-docx-Standards (US-Englisch)
    -- sonst verwirren Rechtschreibprüfung und Silbentrennung beim
    Weiterbearbeiten.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    document.core_properties.language = lang
    normal_rpr = document.styles["Normal"].element.get_or_add_rPr()
    lang_element = normal_rpr.find(qn("w:lang"))
    if lang_element is None:
        lang_element = OxmlElement("w:lang")
        normal_rpr.append(lang_element)
    lang_element.set(qn("w:val"), lang)
    lang_element.set(qn("w:eastAsia"), lang)


def _write_lines(paragraph, text, *, size=None, bold=False):
    """Mehrzeiliger Text als Zeilenumbrüche (`<w:br/>`) innerhalb *eines*
    Absatzes -- ein Absatz je Zeile würde eine Adresse/Fußzeile in mehrere,
    unabhängig umbrechende Absätze zerlegen.
    """
    from docx.shared import Pt

    lines = (text or "").split("\n")
    for index, line in enumerate(lines):
        run = paragraph.add_run(line)
        run.bold = bold
        if size is not None:
            run.font.size = Pt(size)
        if index < len(lines) - 1:
            run.add_break()


def _write_footer(document, content: LetterContent):
    """Dezente Fußzeile mit der eigenen Identität (#1151): dünne Trennlinie
    darüber, kleiner Schriftgrad, auf jeder Seite -- deshalb sowohl in
    `section.footer` als auch in `section.first_page_footer` geschrieben.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if not content.sender_footer_lines:
        return

    section = document.sections[0]
    section.different_first_page_header_footer = True

    for footer in (section.footer, section.first_page_footer):
        footer.is_linked_to_previous = False
        for index, line in enumerate(content.sender_footer_lines):
            item = footer.paragraphs[0] if index == 0 else footer.add_paragraph()
            item.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if index == 0:
                _add_border(item, side="top", size_pt=0.5, color="999999", space_pt=4)
            _write_lines(item, line, size=_FOOTER_FONT_SIZE)


def _write_continuation_header(document, content: LetterContent, printable_width_mm):
    """Verkleinerter Kopfbereich ab Seite 2: Betreff links, Seitenzahl
    rechts -- Seite 1 zeigt keinen Kopf, dort steht die Anschrift bereits
    im Rahmen.
    """
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.shared import Mm, Pt

    section = document.sections[0]
    section.different_first_page_header_footer = True
    header = section.header
    header.is_linked_to_previous = False

    label = content.subject.strip() or content.sender_line
    item = header.paragraphs[0]
    item.paragraph_format.tab_stops.add_tab_stop(Mm(printable_width_mm), WD_TAB_ALIGNMENT.RIGHT)
    _write_lines(item, label, size=_SMALL_FONT_SIZE)
    tab_run = item.add_run("\t")
    tab_run.font.size = Pt(_SMALL_FONT_SIZE)
    prefix_run = item.add_run("Seite ")
    prefix_run.font.size = Pt(_SMALL_FONT_SIZE)
    _add_page_number_field(item, size_pt=_SMALL_FONT_SIZE)
    _add_border(item, side="bottom", size_pt=0.5, color="999999", space_pt=4)


def render_docx(content: LetterContent) -> bytes:
    """Word-Master (.docx) über python-docx, DIN 5008 Form B.

    Word ist das *editierbare* Artefakt: was hier entsteht, soll jemand
    ohne Findus weiterschreiben können. Anschriftfeld, Informationsblock und
    Falz-/Lochmarken sind deshalb Textrahmen an fester Position statt
    gestapelter Leerzeilen (siehe Modul-Docstring oben) -- nur so bleibt die
    Anschrift exakt im Sichtfenster eines Fensterumschlags DIN lang, egal
    wie hoch Briefkopf oder Anschrift ausfallen.
    """
    from docx import Document as DocxDocument
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Mm, Pt

    document = DocxDocument()
    _set_document_language(document)

    section = document.sections[0]
    section.page_width = Mm(PAGE_WIDTH_MM)
    section.page_height = Mm(PAGE_HEIGHT_MM)
    section.left_margin = Mm(PAGE_MARGIN_LEFT_MM)
    section.right_margin = Mm(PAGE_MARGIN_RIGHT_MM)
    section.top_margin = Mm(PAGE_MARGIN_TOP_MM)
    section.bottom_margin = Mm(PAGE_MARGIN_BOTTOM_MM)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(_BODY_FONT_SIZE)

    printable_width_mm = PAGE_WIDTH_MM - PAGE_MARGIN_LEFT_MM - PAGE_MARGIN_RIGHT_MM

    # -- Kopfzone (0-45 mm): Logo rechts, Briefkopftext links, beide als
    # Rahmen -- sie sollen nicht mitzählen, wie weit der Fließtext später
    # nach unten geschoben wird.
    if content.logo:
        logo_paragraph = document.add_paragraph()
        _pin_to_page(
            logo_paragraph,
            x_mm=PAGE_WIDTH_MM - PAGE_MARGIN_RIGHT_MM - LOGO_WIDTH_MM,
            y_mm=10,
            width_mm=LOGO_WIDTH_MM,
        )
        try:
            logo_paragraph.add_run().add_picture(BytesIO(content.logo), width=Mm(LOGO_WIDTH_MM))
        except Exception:
            # Ein Bildformat, das python-docx nicht kennt, darf den Brief
            # nicht kosten -- der Rahmen bleibt dann eben leer.
            logger.warning("Brief-Rendering: Logo konnte nicht eingebettet werden", exc_info=True)

    if content.letterhead.strip():
        letterhead_paragraph = document.add_paragraph()
        letterhead_width = printable_width_mm - (LOGO_WIDTH_MM + 5 if content.logo else 0)
        _pin_to_page(
            letterhead_paragraph, x_mm=PAGE_MARGIN_LEFT_MM, y_mm=10, width_mm=letterhead_width
        )
        _write_lines(letterhead_paragraph, content.letterhead.strip(), bold=True)

    # -- Anschriftfeld: Rücksendeangabe (Zusatz-/Vermerkzone) und Anschrift
    # (Anschriftzone) im selben Rahmen -- ein Rahmen mit zwei Absätzen
    # bleibt trotzdem *ein* Textrahmen, solange beide Absätze dieselbe
    # Position tragen.
    if content.sender_line or content.recipient_block.strip():
        if content.sender_line:
            sender_paragraph = document.add_paragraph()
            _write_lines(sender_paragraph, content.sender_line, size=_SMALL_FONT_SIZE)
            _pin_to_page(
                sender_paragraph,
                x_mm=PAGE_MARGIN_LEFT_MM,
                y_mm=ADDRESS_FIELD_TOP_MM,
                width_mm=ADDRESS_FIELD_WIDTH_MM,
                height_mm=ADDRESS_FIELD_HEIGHT_MM,
            )
        if content.recipient_block.strip():
            recipient_paragraph = document.add_paragraph()
            if content.sender_line:
                recipient_paragraph.paragraph_format.space_before = Pt(6)
            _write_lines(recipient_paragraph, content.recipient_block.strip())
            _pin_to_page(
                recipient_paragraph,
                x_mm=PAGE_MARGIN_LEFT_MM,
                y_mm=ADDRESS_FIELD_TOP_MM,
                width_mm=ADDRESS_FIELD_WIDTH_MM,
                height_mm=ADDRESS_FIELD_HEIGHT_MM,
            )

    # -- Informationsblock rechts: Datum auf Höhe der Anschriftzone.
    if content.date_line:
        date_paragraph = document.add_paragraph()
        date_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _write_lines(date_paragraph, content.date_line)
        _pin_to_page(
            date_paragraph,
            x_mm=PAGE_WIDTH_MM - PAGE_MARGIN_RIGHT_MM - INFO_BLOCK_WIDTH_MM,
            y_mm=INFO_BLOCK_TOP_MM,
            width_mm=INFO_BLOCK_WIDTH_MM,
        )

    # -- Falz- und Lochmarken.
    for y_mm in (FOLD_MARK_1_MM, PUNCH_MARK_MM, FOLD_MARK_2_MM):
        mark_paragraph = document.add_paragraph()
        _pin_to_page(
            mark_paragraph,
            x_mm=0,
            y_mm=y_mm,
            width_mm=_REGISTRATION_MARK_WIDTH_MM,
            height_mm=1,
            height_rule="exact",
        )
        _add_border(mark_paragraph, side="bottom", size_pt=0.75, color="000000", space_pt=0)

    # -- Fließtext ab der Betreffzeile: der einzige Teil, der normal
    # fließt und sich über mehrere Seiten fortsetzen darf. Die erste
    # fließende Zeile bekommt den kompletten Vorschub bis zur Betreffzeile
    # als `space_before` -- die Rahmen oben zählen dafür nicht mit.
    first_flow = True

    def flow_paragraph(text="", *, size=None, bold=False, space_after=6):
        nonlocal first_flow
        item = document.add_paragraph()
        item.paragraph_format.space_after = Pt(space_after)
        if first_flow:
            item.paragraph_format.space_before = Mm(SUBJECT_LINE_TOP_MM - PAGE_MARGIN_TOP_MM)
            first_flow = False
        _write_lines(item, text, size=size, bold=bold)
        return item

    if content.subject.strip():
        flow_paragraph(content.subject.strip(), bold=True, space_after=12)

    for block in re.split(r"\n\s*\n", (content.body or "").strip()):
        if block.strip():
            flow_paragraph(block.strip(), space_after=10)

    if content.closing:
        flow_paragraph(content.closing, space_after=6)

    if content.signature.strip():
        # Unterschriftszone: drei Leerzeilen für die handschriftliche
        # Unterschrift, danach der Name in Klartext.
        for _ in range(3):
            flow_paragraph("", space_after=0)
        flow_paragraph(content.signature.strip(), space_after=0)

    _write_footer(document, content)
    _write_continuation_header(document, content, printable_width_mm)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


# -- PDF -------------------------------------------------------------------


def _pdf_safe(text: str) -> str:
    """fpdf2-Kernschriften können nur cp1252 -- nicht abbildbare Zeichen
    (Emoji, nicht-lateinische Schrift) durch '?' ersetzen, statt beim
    PDF-Bau zu crashen. Umlaute, Euro und typografische Zeichen bleiben.

    Dieselbe Einschränkung wie bei den Mail-Body-PDFs (#1070); bewusst hier
    noch einmal formuliert, statt `apps.ingest` etwas zu entleihen --
    `apps.ingest` hängt von `apps.documents` ab, nicht umgekehrt.
    """
    return (text or "").encode("cp1252", errors="replace").decode("cp1252")


def render_pdf(content: LetterContent) -> bytes:
    """Druck-/Sendefassung (PDF) über fpdf2, direkt aus `LetterContent`.

    Positioniert Absenderzeile und Anschriftfeld absolut (DIN-5008-nah),
    danach läuft der Text im normalen Fluss mit automatischem Seitenumbruch.
    """
    from fpdf import FPDF
    from fpdf.enums import XPos, YPos

    pdf = FPDF(format="A4")
    pdf.core_fonts_encoding = "cp1252"
    pdf.set_margins(PAGE_MARGIN_LEFT_MM, PAGE_MARGIN_TOP_MM, PAGE_MARGIN_RIGHT_MM)
    pdf.set_auto_page_break(auto=True, margin=PAGE_MARGIN_BOTTOM_MM)
    pdf.add_page()

    def line(text, *, height=5.5):
        pdf.multi_cell(0, height, _pdf_safe(text) or " ", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    if content.logo:
        try:
            pdf.image(
                BytesIO(content.logo),
                x=pdf.w - PAGE_MARGIN_RIGHT_MM - LOGO_WIDTH_MM,
                y=PAGE_MARGIN_TOP_MM,
                w=LOGO_WIDTH_MM,
            )
        except Exception:
            logger.warning("Brief-Rendering: Logo konnte nicht ins PDF gesetzt werden", exc_info=True)

    if content.letterhead.strip():
        pdf.set_font("Helvetica", style="B", size=12)
        line(content.letterhead.strip(), height=6)

    # Anschriftfeld: feste Oberkante, damit der Empfänger im Fensterumschlag
    # sichtbar ist -- unabhängig davon, wie hoch der Briefkopf ausgefallen ist.
    pdf.set_y(max(pdf.get_y(), ADDRESS_FIELD_TOP_MM - 8))
    if content.sender_line:
        pdf.set_font("Helvetica", size=_SMALL_FONT_SIZE)
        line(content.sender_line, height=4)

    pdf.set_y(max(pdf.get_y(), ADDRESS_FIELD_TOP_MM))
    pdf.set_font("Helvetica", size=_BODY_FONT_SIZE)
    if content.recipient_block.strip():
        line(content.recipient_block.strip())

    if content.date_line:
        pdf.ln(8)
        pdf.cell(0, 5.5, _pdf_safe(content.date_line), align="R", new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.ln(8)
    if content.subject.strip():
        pdf.set_font("Helvetica", style="B", size=_BODY_FONT_SIZE)
        line(content.subject.strip(), height=6)
        pdf.ln(4)

    pdf.set_font("Helvetica", size=_BODY_FONT_SIZE)
    for block in re.split(r"\n\s*\n", (content.body or "").strip()):
        if not block.strip():
            continue
        line(block.strip())
        pdf.ln(3)

    if content.closing:
        pdf.ln(4)
        line(content.closing)

    if content.signature.strip():
        pdf.ln(8)
        line(content.signature.strip())

    output = pdf.output()
    return bytes(output)


# -- Ablage am Entwurf -----------------------------------------------------


def file_slug(content: LetterContent, draft) -> str:
    """Dateiname-Stamm für beide Ausgaben: „brief-<betreff>-<id>".

    Die Entwurfs-ID hängt dran, damit zwei Briefe mit demselben Betreff im
    Downloads-Ordner unterscheidbar bleiben.
    """
    slug = re.sub(r"[^\w\- ]", "", content.subject or "").strip().replace(" ", "-")[:60]
    return f"brief-{slug}-{draft.pk}".replace("--", "-").strip("-").lower()


def render_draft_files(draft) -> LetterContent:
    """Rendert Word + PDF und legt beide am Entwurf ab.

    Ersetzt vorhandene Dateien bei jedem Aufruf -- „Text ändern und neu
    rendern" darf keine Halde alter Fassungen im Object-Storage
    hinterlassen, und die einzige Fassung, die zählt, ist die, die der
    Nutzer gerade vor sich sieht.
    """
    from django.core.files.base import ContentFile

    content = build_content(draft)
    slug = file_slug(content, draft)

    for field_name, suffix, data in (
        ("docx_file", "docx", render_docx(content)),
        ("pdf_file", "pdf", render_pdf(content)),
    ):
        file_field = getattr(draft, field_name)
        if file_field:
            file_field.delete(save=False)
        file_field.save(f"{slug}.{suffix}", ContentFile(data), save=False)

    draft.save(update_fields=["docx_file", "pdf_file", "updated_at"])
    return content
